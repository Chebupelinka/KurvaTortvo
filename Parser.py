from openpyxl import load_workbook
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import os


class Parser:

    nameOfList = "Лист1"

    # Здесь массив букв, до куда таблица идет
    LettersInBook = ['Q', 'R', 'S', 'T', 'U', 'V', 'W', 'X', 'Y', 'Z', 'AA', 'AB', 'AC', 'AD', 'AE', 'AF', 'AG',
                     'AH', 'AI', 'AJ']

    LettersForSmeshariki = ['AK', 'AL', 'AM', 'AN']

    # Список факультетов
    faculties = []
    path = ""
    book = None
    sheet = None
    i = 2

    def __init__(self, new_path, book_path, new_faculties=["ФРТ", "ФЭЛ", "ФКТИ", "ФЭА", "ФИБС", "ИНПРОТЕХ", "ГФ"]):
        self.path = new_path
        self.faculties = new_faculties
        self.book = load_workbook(filename=book_path)
        self.sheet = self.book[self.nameOfList]

    def creating_dirs(self, root_id, ):

        for folder in self.faculties:



while str(sheet['C' + str(i)].value) != "None":
    index = faculties.index(sheet['F' + str(i)].value)
    new_path = path + "\\" + faculties[index]
    file_name = "Очное собеседование - " + sheet['C' + str(i)].value + ' ' + sheet['D' + str(i)].value + ".docx"
    new_path_word = new_path + "\\" + file_name
    if os.path.exists(new_path_word):
        i += 1
        continue
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)
    parag = doc.add_paragraph()

    print(i)

    if str(sheet['E' + str(i)].value) != "None":
        run = parag.add_run(sheet['C' + str(i)].value + ' ' + sheet['D' + str(i)].value + ' ' + sheet['E' + str(i)].value +
                        ' ' + str(sheet['H' + str(i)].value) + ' ' + sheet['I' + str(i)].value + ' ' +
                        str(sheet['M' + str(i)].value))
    else:
        run = parag.add_run(sheet['C' + str(i)].value + ' ' + sheet['D' + str(i)].value + ' ' +
                        str(sheet['H' + str(i)].value) + ' ' + sheet['I' + str(i)].value + ' ' +
                        str(sheet['M' + str(i)].value))

    parag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font = run.font
    font.name = 'Arial'
    font.size = Pt(10)

    if str(sheet['O' + str(i)].value) != "None":
        table = doc.add_table(rows=len(LettersInBook), cols=2, style='Table Grid')

        j = 0
        while j < len(LettersInBook):
            left = table.cell(j, 0)
            left.width = Cm(5.5)
            left.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            left_p = left.add_paragraph()
            left_p.text = ""

            right = table.cell(j, 1)
            right.width = Cm(14)
            right.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            right_p = right.add_paragraph()

            left_run = left_p.add_run(sheet[LettersInBook[j] + '1'].value.rstrip())

            try:
                right_run = right_p.add_run(sheet[LettersInBook[j] + str(i)].value.rstrip())
            except Exception:
                right_run = right_p.add_run("Ничего не написали(((".rstrip())

            left_font = left_run.font
            left_font.name = 'Arial'
            left_font.size = Pt(10)

            right_font = right_run.font
            right_font.name = 'Arial'
            right_font.size = Pt(10)

            j += 1
    else:
        table = doc.add_table(rows=len(LettersForSmeshariki), cols=2, style='Table Grid')

        j = 0
        while j < len(LettersForSmeshariki):
            left = table.cell(j, 0)
            left.width = Cm(5.5)
            left.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            left_p = left.add_paragraph()

            right = table.cell(j, 1)
            right.width = Cm(14)
            right.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            right_p = right.add_paragraph()

            left_run = left_p.add_run(sheet[LettersForSmeshariki[j] + '1'].value.rstrip())

            try:
                right_run = right_p.add_run(sheet[LettersForSmeshariki[j] + str(i)].value.rstrip())
            except Exception:
                right_run = right_p.add_run("Ничего не написали(((".rstrip())

            left_font = left_run.font
            left_font.name = 'Arial'
            left_font.size = Pt(10)

            right_font = right_run.font
            right_font.name = 'Arial'
            right_font.size = Pt(10)

            j += 1
    doc.save(new_path_word)
    i += 1
